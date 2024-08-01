import os
from docx import Document
from markdownify import markdownify as md

def convert_table_to_markdown(table):
   """Convert a docx table to a Markdown table."""
   rows = []
   for row in table.rows:
       cells = [cell.text.strip() for cell in row.cells]
       rows.append("| " + " | ".join(cells) + " |")


   # Create the Markdown table header separator
   if rows:
       num_columns = len(rows[0].split("|")) - 2  # Subtract 2 for the leading and trailing '|'
       header_separator = "| " + " | ".join(["---"] * num_columns) + " |"
       rows.insert(1, header_separator)


   return "\n".join(rows)

def convert_docx_to_markdown(input_path, output_path):
   try:
       # Load the .docx file
       doc = Document(input_path)
       # Extract the text and convert to HTML
       html_content = ""
       for para in doc.paragraphs:
           html_content += para.text + "<br>"


       # Extract tables and convert to Markdown
       for table in doc.tables:
           table_md = convert_table_to_markdown(table)
           html_content += table_md + "<br>"


       # Debugging: Print the HTML content before conversion
       print("HTML Content Before Conversion:")
       print(html_content)


       # Convert HTML to Markdown
       markdown_content = md(html_content)


       # Debugging: Print the Markdown content after conversion
       print("Markdown Content After Conversion:")
       print(markdown_content)


       # Save the Markdown content to a file with utf-8 encoding
       with open(output_path, 'w', encoding='utf-8') as f:
           f.write(markdown_content)


       print(f"Converted {input_path} to {output_path}")
   except Exception as e:
       print(f"Failed to convert {input_path}: {e}")


def convert_all_docx_in_directory(input_directory, output_directory):
   if not os.path.exists(output_directory):
       os.makedirs(output_directory)


   for filename in os.listdir(input_directory):
       if filename.endswith(".docx"):
           input_path = os.path.join(input_directory, filename)
           output_filename = os.path.splitext(filename)[0] + ".md"
           output_path = os.path.join(output_directory, output_filename)
           convert_docx_to_markdown(input_path, output_path)


if __name__ == "__main__":
   input_directory = "C:/Users/codeineteh/Downloads/Policy Docx/Policy Docx"
   output_directory = "C:/Users/codeineteh/PycharmProjects/pythonProject/masterDB"
   convert_all_docx_in_directory(input_directory, output_directory)

